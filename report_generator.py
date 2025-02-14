"""
Word文档报告生成器

该模块负责生成格式化的Word文档报告，包括：
1. 统一的字体和样式设置
2. 项目信息的结构化展示
3. 分析结果的格式化处理
4. 文档布局的优化

使用示例：
    report = ReportGenerator()
    report.add_project(project_info, analysis_text)
    report.save("report.docx")
"""

from docx import Document
import docx.opc.constants
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import qn as shared_qn
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import logging
import os

logger = logging.getLogger(__name__)

class ReportGenerator:
    """
    Word文档报告生成器类
    
    负责创建和管理Word文档，设置统一的样式，
    添加项目信息和分析结果，并保存为格式化的文档。
    
    属性:
        document: Document对象，用于操作Word文档
    """

    def __init__(self):
        """
        初始化报告生成器
        
        创建新的Word文档，设置默认字体和样式，
        添加文档标题和时间戳。
        """
        self.document = Document()
        # 设置默认字体
        self._set_default_font()
        self._create_title()
        logger.debug("初始化报告生成器")

    def _set_default_font(self):
        """
        设置文档默认字体
        
        配置全局字体样式：
        - 字体：微软雅黑
        - 字号：12pt
        - 确保中文字体正确显示
        """
        # 设置默认字体
        self.document.styles['Normal'].font.name = '微软雅黑'
        self.document.styles['Normal'].font.size = Pt(12)
        # 设置中文字体
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _create_title(self):
        """
        创建报告标题
        
        添加文档主标题和时间戳：
        - 标题：20pt，加粗，居中
        - 时间：12pt，居中
        - 统一使用微软雅黑字体
        """
        # 添加标题
        title = self.document.add_heading('GitHub优质项目分析报告', level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 设置标题格式
        run = title.runs[0]
        run.font.name = '微软雅黑'
        run.font.size = Pt(20)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 添加生成时间
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M')
        time_paragraph = self.document.add_paragraph()
        time_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        time_run = time_paragraph.add_run(f'生成时间：{current_time}')
        time_run.font.size = Pt(12)
        time_run.font.name = '微软雅黑'
        time_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _create_element(self, name):
        """创建OxmlElement元素"""
        return OxmlElement(name)

    def _create_hyperlink(self, paragraph, url, text):
        """
        创建超链接
        
        Args:
            paragraph: 段落对象
            url: 链接地址
            text: 显示文本
        """
        # 创建hyperlink标签
        hyperlink = self._create_element('w:hyperlink')
        
        # 创建链接关系
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink.set(qn('r:id'), r_id)
        
        # 创建文本运行对象
        new_run = self._create_element('w:r')
        rPr = self._create_element('w:rPr')
        
        # 设置样式
        rStyle = self._create_element('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        
        # 设置字体
        rFonts = self._create_element('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
        rPr.append(rFonts)
        
        # 设置字号
        sz = self._create_element('w:sz')
        sz.set(qn('w:val'), '24')
        rPr.append(sz)
        
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        
        paragraph._p.append(hyperlink)
        return hyperlink

    def add_project(self, project: dict, analysis: str):
        """
        添加项目分析内容
        
        Args:
            project (dict): 项目信息字典，包含：
                - name: 项目名称
                - url: 项目地址
                - stars: star数量
                - forks: fork数量
                - language: 主要编程语言
                - description: 项目描述
            analysis (str): AI分析结果文本
        
        格式说明：
        - 项目标题：16pt，加粗
        - 项目信息：12pt，普通
        - URL：12pt，超链接
        - 分析内容：12pt，普通
        - 分隔线：灰色，居中
        """
        # 添加项目标题（加粗，16号字）
        heading = self.document.add_heading('', level=1)
        title_run = heading.add_run(project['name'])
        title_run.font.name = '微软雅黑'
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 添加项目基本信息
        info_paragraph = self.document.add_paragraph()
        
        # 添加项目地址（超链接）
        url_run = info_paragraph.add_run("项目地址：")
        url_run.font.name = '微软雅黑'
        url_run.font.size = Pt(12)
        url_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 使用新的方法添加超链接
        self._create_hyperlink(info_paragraph, project['url'], project['url'])
        info_paragraph.add_run('\n')
        
        # 添加其他信息
        info_text = [
            f"Stars数量：{project['stars']}",
            f"Fork数量：{project['forks']}",
            f"主要语言：{project.get('language', '未知')}",
            f"项目描述：{project.get('description', '无描述')}"
        ]
        
        for text in info_text:
            run = info_paragraph.add_run(text + '\n')
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 添加分析内容
        analysis_paragraph = self.document.add_paragraph()
        analysis_run = analysis_paragraph.add_run(analysis)
        analysis_run.font.name = '微软雅黑'
        analysis_run.font.size = Pt(12)
        analysis_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 添加分隔线
        separator = self.document.add_paragraph()
        separator.paragraph_format.space_after = Pt(20)
        separator.paragraph_format.space_before = Pt(20)
        separator_run = separator.add_run('=' * 50)
        separator_run.font.size = Pt(12)
        separator_run.font.color.rgb = RGBColor(128, 128, 128)

        logger.info(f"📄 生成报告 | 项目: {project['name']}")

    def save(self, filename: str):
        """
        保存报告文件
        
        Args:
            filename (str): 保存的文件路径
        
        页面设置：
        - 页边距：1英寸
        - 支持中文字符
        - 自动调整布局
        """
        # 设置页面边距
        sections = self.document.sections
        for section in sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

        self.document.save(filename)
        logger.info(f"✅ 报告已保存：{filename}") 